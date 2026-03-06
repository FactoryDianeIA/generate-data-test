import json
import os
import fitz  # PyMuPDF pour les PDF
from docx import Document

def extract_text_universally(file_path, output_json):
    """Extrait le texte de n'importe quel format (incluant les tableaux Word)."""
    extension = os.path.splitext(file_path)[1].lower()
    full_text = []

    if extension == ".docx":
        doc = Document(file_path)
        
        # 1. Extraction des paragraphes
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 2. EXTRACTION DES TABLEAUX (C'est ici que se trouvent vos 17 colonnes !)
        for table in doc.tables:
            for row in table.rows:
                # On concatène les cellules d'une ligne avec un séparateur |
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    full_text.append(" | ".join(cells))
        
        content = "\n".join(full_text)
    
    elif extension == ".pdf":
        with fitz.open(file_path) as doc:
            content = "\n".join([page.get_text() for page in doc])
            
    elif extension in [".md", ".txt"]:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            
    else:
        raise ValueError(f"Format de fichier non supporté : {extension}")

    # Sauvegarde pour l'agent Gemini
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump({"filename": os.path.basename(file_path), "text": content}, f, indent=4)